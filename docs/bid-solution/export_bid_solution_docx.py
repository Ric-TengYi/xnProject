from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "exported-docx"

EXPORTS = [
    {
        "title": "渣土运输与消纳监管平台需求方案",
        "subject": "甲方版需求方案",
        "output": OUTPUT_DIR / "甲方版-需求方案.docx",
        "files": sorted((ROOT / "01-甲方版-需求方案").glob("*.md")),
    },
    {
        "title": "渣土运输与消纳监管平台技术方案",
        "subject": "甲方版技术方案",
        "output": OUTPUT_DIR / "甲方版-技术方案.docx",
        "files": sorted((ROOT / "02-甲方版-技术方案").glob("*.md")),
    },
    {
        "title": "渣土运输与消纳监管平台专项方案",
        "subject": "甲方版专项方案",
        "output": OUTPUT_DIR / "甲方版-专项方案.docx",
        "files": sorted((ROOT / "03-专项方案").glob("*.md")),
    },
    {
        "title": "渣土运输与消纳监管平台商务技术文件章节",
        "subject": "商务技术文件第5章至第10章",
        "output": OUTPUT_DIR / "商务技术文件章节-汇编.docx",
        "files": sorted(
            file
            for file in (ROOT / "05-商务技术文件章节").glob("[0-9][0-9]_*.md")
            if file.stem[:2] >= "05"
        ),
    },
]

INDIVIDUAL_EXPORT_DIRS = [
    ROOT / "05-商务技术文件章节",
]

EXTRA_SINGLE_EXPORTS = [
    ROOT / "投标交付检查清单.md",
]


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^-\s+(.*)$")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)$")


def configure_document(document: Document, title: str, subject: str) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name, size, bold in [("Heading 1", 18, True), ("Heading 2", 16, True), ("Heading 3", 14, True)]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(9)

    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    core = document.core_properties
    core.title = title
    core.subject = subject
    core.author = "项目投标方案编制组"
    core.company = "项目投标方案编制组"


def add_cover(document: Document, title: str, subject: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(18)
    run2 = p2.add_run(subject)
    run2.font.size = Pt(14)
    run2.font.name = "Microsoft YaHei"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(260)
    run3 = p3.add_run(f"导出日期：{date.today().isoformat()}")
    run3.font.size = Pt(12)
    run3.font.name = "Microsoft YaHei"
    run3._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    document.add_page_break()


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    for idx, line in enumerate(clean_inline_text(text).split("\n")):
        if idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def clean_inline_text(text: str) -> str:
    text = text.replace("<br/>", "\n").replace("<br>", "\n")
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    return text.strip()


def insert_bookmarkless_toc_hint(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("目录")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    toc = document.add_paragraph()
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    toc._p.append(fld_simple)
    hint = document.add_paragraph("说明：如目录未自动刷新，可在 Word 中全选后更新域。")
    hint.paragraph_format.space_after = Pt(8)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    fld_simple.append(r)
    p._p.append(fld_simple)


def render_markdown(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            document.add_heading(clean_inline_text(heading_match.group(2)), level=level)
            i += 1
            continue

        image_match = IMAGE_RE.match(line.strip())
        if image_match:
            image_path = (markdown_path.parent / image_match.group(2)).resolve()
            if image_path.exists():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(image_path), width=Cm(15.5))
            else:
                document.add_paragraph(f"[缺失图片] {image_match.group(2)}")
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(document, table_lines)
            continue

        ordered_match = ORDERED_RE.match(line.strip())
        if ordered_match:
            while i < len(lines):
                match = ORDERED_RE.match(lines[i].strip())
                if not match:
                    break
                p = document.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                p.add_run(clean_inline_text(match.group(2)))
                i += 1
            continue

        unordered_match = UNORDERED_RE.match(line.strip())
        if unordered_match:
            while i < len(lines):
                match = UNORDERED_RE.match(lines[i].strip())
                if not match:
                    break
                p = document.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.add_run(clean_inline_text(match.group(1)))
                i += 1
            continue

        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            document.add_paragraph("\n".join(code_lines), style="CodeBlock")
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i].rstrip()
            if not next_line.strip():
                break
            if (
                HEADING_RE.match(next_line)
                or next_line.strip().startswith("|")
                or ORDERED_RE.match(next_line.strip())
                or UNORDERED_RE.match(next_line.strip())
                or IMAGE_RE.match(next_line.strip())
                or next_line.strip().startswith("```")
            ):
                break
            paragraph_lines.append(next_line)
            i += 1
        p = document.add_paragraph(clean_inline_text(" ".join(paragraph_lines)))
        p.paragraph_format.space_after = Pt(6)


def add_markdown_table(document: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        document.add_paragraph(clean_inline_text(" ".join(table_lines)))
        return

    rows = []
    for line in table_lines:
        parts = [clean_inline_text(part) for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    header = rows[0]
    body = [row for row in rows[2:] if row]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True
    for idx, value in enumerate(header):
        set_cell_text(table.rows[0].cells[idx], value)
    for row in body:
        cells = table.add_row().cells
        for idx in range(len(header)):
            set_cell_text(cells[idx], row[idx] if idx < len(row) else "")
    document.add_paragraph()


def build_docx(export: dict) -> Path:
    document = Document()
    configure_document(document, export["title"], export["subject"])
    add_cover(document, export["title"], export["subject"])
    insert_bookmarkless_toc_hint(document)
    for idx, markdown_file in enumerate(export["files"]):
        render_markdown(document, markdown_file)
        if idx < len(export["files"]) - 1:
            document.add_section(WD_SECTION_START.NEW_PAGE)
    for section in document.sections:
        add_page_number(section)
    export["output"].parent.mkdir(parents=True, exist_ok=True)
    document.save(export["output"])
    return export["output"]


def extract_title(markdown_path: Path) -> str:
    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        match = HEADING_RE.match(line.strip())
        if match:
            return clean_inline_text(match.group(2))
    return markdown_path.stem


def build_single_markdown_docx(markdown_path: Path, output_path: Path) -> Path:
    title = extract_title(markdown_path)
    document = Document()
    configure_document(document, title, title)
    add_cover(document, title, "商务技术文件章节")
    render_markdown(document, markdown_path)
    for section in document.sections:
        add_page_number(section)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def main() -> None:
    generated = [build_docx(export) for export in EXPORTS]
    for directory in INDIVIDUAL_EXPORT_DIRS:
        target_dir = OUTPUT_DIR / directory.name
        for markdown_file in sorted(
            file for file in directory.glob("[0-9][0-9]_*.md") if file.stem[:2] >= "05"
        ):
            generated.append(build_single_markdown_docx(markdown_file, target_dir / f"{markdown_file.stem}.docx"))
    for markdown_file in EXTRA_SINGLE_EXPORTS:
        generated.append(build_single_markdown_docx(markdown_file, OUTPUT_DIR / f"{markdown_file.stem}.docx"))
    for path in generated:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
